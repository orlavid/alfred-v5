"""DOCX output writer."""

from __future__ import annotations

from pathlib import Path
from typing import Dict

from docx import Document


class DocxWriter:
    def __init__(self, output_dir: Path) -> None:
        self.output_dir = output_dir
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def write(self, handbook: Dict[str, str]) -> None:
        doc = Document()
        doc.add_heading("Alfred Engineering Handbook", 0)
        doc.add_paragraph("Recovery Point Alpha")

        for filename, markdown in handbook.items():
            doc.add_page_break()
            doc.add_heading(filename.replace("-", " ").replace(".md", "").title(), 1)

            in_code = False
            for line in markdown.splitlines():
                stripped = line.strip()
                if stripped.startswith("```"):
                    in_code = not in_code
                    continue
                if not stripped:
                    continue
                if in_code:
                    doc.add_paragraph(stripped)
                elif stripped.startswith("# "):
                    doc.add_heading(stripped[2:], 1)
                elif stripped.startswith("## "):
                    doc.add_heading(stripped[3:], 2)
                elif stripped.startswith("### "):
                    doc.add_heading(stripped[4:], 3)
                elif stripped.startswith("- "):
                    doc.add_paragraph(stripped[2:], style="List Bullet")
                else:
                    doc.add_paragraph(stripped)

        doc.save(self.output_dir / "Alfred_Engineering_Handbook_Recovery_Point_Alpha.docx")
